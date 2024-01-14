import os
import streamlit as st
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# OCR function to extract text from images
def extract_text_from_image(image_path):
    img = Image.open(image_path)
    text = pytesseract.image_to_string(img)
    return text.strip() if text else None

# Function to extract images from a Word document
def extract_images_from_word(docx_path):
    document = Document(docx_path)
    images = []

    for rel in document.part.rels:
        if "image" in document.part.rels[rel].target_ref:
            image_part = document.part.rels[rel].target_part
            image_data = image_part.blob
            images.append(image_data)

    return images

# Streamlit App
st.title("Image and Text Viewer with OCR")

# File upload
uploaded_file = st.file_uploader("Upload a Word Document", type=["docx"])

if uploaded_file:
    # Temporary folder for storing images
    image_folder = 'temp_images'
    os.makedirs(image_folder, exist_ok=True)

    # Save the uploaded Word document to a temporary file
    word_doc_path = os.path.join(image_folder, "uploaded.docx")
    with open(word_doc_path, "wb") as f:
        f.write(uploaded_file.getvalue())

    # Extract images from the Word document
    images = extract_images_from_word(word_doc_path)

    # Display only images with extracted text
    for i, image_data in enumerate(images):
        image_path = os.path.join(image_folder, f"image_{i + 1}.png")
        with open(image_path, "wb") as img_file:
            img_file.write(image_data)

        # Extract text from the image using OCR
        extracted_text = extract_text_from_image(image_path)

        if extracted_text:
            st.subheader(f"Image {i + 1}")
            st.image(image_data, caption=f"Image")

            st.text_area(f"Extracted Text from Image {i + 1}", extracted_text, height=300)  # Set the height parameter here
        else:
            st.info(f"No text extracted from Image {i + 1}")

    # Download Parsed Text as Word
    download_button = st.button("Download Parsed Text as Word")
    if download_button:
        # Save extracted text to a Word document
        doc = Document()
        for i, image_data in enumerate(images):
            image_path = os.path.join(image_folder, f"image_{i + 1}.png")
            extracted_text = extract_text_from_image(image_path)
            if extracted_text:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(f"Text from Image {i + 1}:\n\n{extracted_text}\n\n")
                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(11)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Save the Word document
        word_file_path = 'extracted_text.docx'
        doc.save(word_file_path)
        st.success("Parsed text downloaded successfully!")

        # Provide download link
        st.markdown(f"### [Download Extracted Text as Word](./{word_file_path})")